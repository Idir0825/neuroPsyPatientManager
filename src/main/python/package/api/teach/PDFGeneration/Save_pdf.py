from pathlib import Path
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert
from PDFGeneration.Graphique import Make_Graph

CHEMIN_VERS_TEMPLATE = "static/TEA-CH.docx"

class Results_as_pdf:

	def __init__(self, patient, results_for_graph, results_for_pdf, saving_location):

		Make_Graph(patient, results_for_graph, saving_location)  # Makes and saves the results graph

		self.SAVING_LOCATION = f"{saving_location}/" \
		                       f"Bilans_TEA-CH/" \
		                       f"{patient.lastName}_{patient.firstName}_{patient.birthDate}"

		self.SAVING_LOCATION_DOCX = f"{self.SAVING_LOCATION}/Resultats.docx"

		self.SAVING_LOCATION_PDF = f"{self.SAVING_LOCATION}/Resultats.pdf"

		if not os.path.exists(self.SAVING_LOCATION):

			try:
				os.makedirs(self.SAVING_LOCATION)
			except OSError:
				print("Creation of the directory %s failed" % self.SAVING_LOCATION)
			else:
				print("Successfully created the directory %s" % self.SAVING_LOCATION)
		else:
			pass

		self.document = Document(CHEMIN_VERS_TEMPLATE)

		self._fill_document(patient, results_for_pdf)

		self._add_graph(patient, saving_location)

		self.document.save(self.SAVING_LOCATION_DOCX)

		convert(self.SAVING_LOCATION_DOCX, self.SAVING_LOCATION_PDF)  # This converts the wordfile to a pdf file

	def _fill_document(self, patient, notes):

		""" Function to fill the docx TEA-CH template """

		# Adding a style
		styles = self.document.styles

		simple = styles.add_style("Simple", WD_STYLE_TYPE.PARAGRAPH)
		simple_font = simple.font
		simple_font.name = 'Bell MT'
		simple_font.size = Pt(9)


		# -- First Table
		self.document.tables[0].row_cells(0)[0].text = patient.lastName  # Adding lastName
		self.document.tables[0].row_cells(0)[1].text = patient.firstName  # Adding firstName
		self.document.tables[0].row_cells(0)[2].text = f"{patient.age[0]} ans, {patient.age[1]} mois"  # Adding age

		for j in range(0, 3):
			for i in range(0, len(self.document.tables[0].row_cells(0)[j].paragraphs)):
				self.document.tables[0].row_cells(0)[j].paragraphs[i].style = self.document.styles['Simple']
				self.document.tables[0].row_cells(0)[j].paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
				self.document.tables[0].row_cells(0)[j].paragraphs[i].paragraph_format.space_before = Pt(6)
				self.document.tables[0].row_cells(0)[j].paragraphs[i].paragraph_format.space_after = Pt(6)

		# -- Date de passation --
		self.document.paragraphs[5].text = f"Date de passation : {patient.createdOn}"  # Adding creationDate
		self.document.paragraphs[5].style = self.document.styles['Simple']
		self.document.paragraphs[5].alignement = WD_PARAGRAPH_ALIGNMENT.RIGHT
		self.document.paragraphs[5].paragraph_format.space_before = Pt(6)
		self.document.paragraphs[5].paragraph_format.space_after = Pt(6)
		for i in range(0, len(self.document.paragraphs[5].runs)):
			# -- Making in bold
			self.document.paragraphs[5].runs[i].font.bold = True

		# -- Filling the second table cells --
		self.document.tables[1].row_cells(1)[4].paragraphs[0].text = f"B={notes['B'].get()}"
		self.document.tables[1].row_cells(1)[4].paragraphs[1].text = f"C={notes['C'].get()}"
		self.document.tables[1].row_cells(1)[4].paragraphs[2].text = f"G={notes['G'].get()}"  # Adding B, C, G grades
		self.document.tables[1].row_cells(1)[5].paragraphs[0].text = f"{notes['Bp'].get()}"
		self.document.tables[1].row_cells(1)[5].paragraphs[1].text = f"{notes['Cp'].get()}"
		self.document.tables[1].row_cells(1)[5].paragraphs[2].text = f"{notes['Gp'].get()}"  # Adding B, C, G percentils
		self.document.tables[1].row_cells(1)[6].paragraphs[0].text = f"{self._interpretation(notes['Bp'].get())}"
		self.document.tables[1].row_cells(1)[6].paragraphs[1].text = f"{self._interpretation(notes['Cp'].get())}"
		self.document.tables[1].row_cells(1)[6].paragraphs[2].text = f"{self._interpretation(notes['Gp'].get())}"  # Adding B, C, G interpretations

		self.document.tables[1].row_cells(2)[4].paragraphs[0].text = f"U={notes['U'].get()}"  # Adding U grade
		self.document.tables[1].row_cells(2)[5].paragraphs[0].text = f"{notes['Up'].get()}"  # Adding U percentil
		self.document.tables[1].row_cells(2)[6].paragraphs[0].text = f"{self._interpretation(notes['Up'].get())}"  # Adding U interpretation

		self.document.tables[1].row_cells(3)[4].paragraphs[0].text = f"H={notes['H'].get()}"  # Adding H grade
		self.document.tables[1].row_cells(3)[5].paragraphs[0].text = f"{notes['Hp'].get()}"  # Adding H percentil
		self.document.tables[1].row_cells(3)[6].paragraphs[0].text = f"{self._interpretation(notes['Hp'].get())}"  # Adding H interpretation

		self.document.tables[1].row_cells(4)[4].paragraphs[0].text = f"BB={notes['BB'].get()}"  # Adding BB grade
		self.document.tables[1].row_cells(4)[5].paragraphs[0].text = f"{notes['BBp'].get()}"  # Adding BB percentil
		self.document.tables[1].row_cells(4)[6].paragraphs[0].text = f"{self._interpretation(notes['BBp'].get())}"  # Adding BB interpretation

		self.document.tables[1].row_cells(5)[4].paragraphs[0].text = f"X={notes['X'].get()}"  # Adding X grade
		self.document.tables[1].row_cells(5)[5].paragraphs[0].text = f"{notes['Xp'].get()}"  # Adding X percentil
		self.document.tables[1].row_cells(5)[6].paragraphs[0].text = f"{self._interpretation(notes['Xp'].get())}"  # Adding U interpretation

		self.document.tables[1].row_cells(6)[4].paragraphs[0].text = f"T={notes['T'].get()}"  # Adding T grade
		self.document.tables[1].row_cells(6)[5].paragraphs[0].text = f"{notes['Tp'].get()}"  # Adding T percentil
		self.document.tables[1].row_cells(6)[6].paragraphs[0].text = f"{self._interpretation(notes['Tp'].get())}"  # Adding T interpretation

		self.document.tables[1].row_cells(7)[4].paragraphs[0].text = f"I={notes['I'].get()}"
		self.document.tables[1].row_cells(7)[4].paragraphs[1].text = f"L={notes['L'].get()}"  # Adding I, L grades
		self.document.tables[1].row_cells(7)[5].paragraphs[0].text = f"{notes['Ip'].get()}"
		self.document.tables[1].row_cells(7)[5].paragraphs[1].text = f"{notes['Lp'].get()}"  # Adding I, L percentils
		self.document.tables[1].row_cells(7)[6].paragraphs[0].text = f"{self._interpretation(notes['Ip'].get())}"
		self.document.tables[1].row_cells(7)[6].paragraphs[1].text = f"{self._interpretation(notes['Lp'].get())}"  # Adding I,
		# L interpretation

		self.document.tables[1].row_cells(8)[4].paragraphs[0].text = f"Z={notes['Z'].get()}"
		self.document.tables[1].row_cells(8)[4].paragraphs[1].text = f"AA={notes['AA'].get()}"  # Adding Z, AA grades
		self.document.tables[1].row_cells(8)[5].paragraphs[0].text = f"{notes['Zp'].get()}"
		self.document.tables[1].row_cells(8)[5].paragraphs[1].text = f"{notes['AAp'].get()}"  # Adding Z, AA percentils
		self.document.tables[1].row_cells(8)[6].paragraphs[0].text = f"{self._interpretation(notes['Zp'].get())}"
		self.document.tables[1].row_cells(8)[6].paragraphs[1].text = f"{self._interpretation(notes['AAp'].get())}"  # Adding Z,
		# AA interpretation

		self.document.tables[1].row_cells(9)[4].paragraphs[0].text = f"Y={notes['Y'].get()}"  # Adding Y grade
		self.document.tables[1].row_cells(9)[5].paragraphs[0].text = f"{notes['Yp'].get()}"  # Adding Y percentil
		self.document.tables[1].row_cells(9)[6].paragraphs[0].text = f"{self._interpretation(notes['Yp'].get())}"  # Adding Y interpretation

		# -- Setting the table paragraphs style and alignement --
		for i in range(1, 10):
			# -- Going through the rows --
			for j in range(4, 7):
				# -- Going through the interesting cells --
				for k in range(0, len(self.document.tables[1].row_cells(i)[j].paragraphs)):
					# Going through the paragraphs of the cell
					self.document.tables[1].row_cells(i)[j].paragraphs[k].style = self.document.styles['Simple']
					self.document.tables[1].row_cells(i)[j].paragraphs[k].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
					self.document.tables[1].row_cells(i)[j].paragraphs[k].paragraph_format.space_before = Pt(6)
					self.document.tables[1].row_cells(i)[j].paragraphs[k].paragraph_format.space_after = Pt(6)

				for p in range(0, len(self.document.tables[1].row_cells(i)[5].paragraphs)):
					# -- Going through the paragraphs of cells 5 and 6 --
					for m in range(0, len(self.document.tables[1].row_cells(i)[5].paragraphs[p].runs)):
						# -- Making the text bold --
						self.document.tables[1].row_cells(i)[5].paragraphs[p].runs[m].font.bold = True
					for n in range(0, len(self.document.tables[1].row_cells(i)[6].paragraphs[p].runs)):
						# -- Making the text italic
						self.document.tables[1].row_cells(i)[6].paragraphs[p].runs[n].font.italic = True

	def _interpretation(self, note):

		if note >= 95:
			interpretation = "Très élevé"
		elif 90 <= note <= 94:
			interpretation = "Élevé"
		elif 75 <= note <= 89:
			interpretation = "Moyen élevé"
		elif 25 <= note <= 74:
			interpretation = "Moyen"
		elif 10 <= note <= 24:
			interpretation = "Moyen faible"
		elif 6 <= note <= 9:
			interpretation = "Faible / Limite"
		else:
			interpretation = "Déficitaire"

		return interpretation

	def _add_graph(self, patient, saving_location):
		self.graph = self.document.add_picture(
			f"{saving_location}/Graphiques/{patient.lastName}_{patient.firstName}_{patient.birthDate}.png",
			width=Inches(7.2)
			)

		self.last_paragraph = self.document.paragraphs[-1]
		self.last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
